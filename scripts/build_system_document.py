from __future__ import annotations

import re
import subprocess
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "generated"
MD_PATH = OUT / "KASIHUB_SYSTEM_REFERENCE.md"
DOCX_PATH = OUT / "KASIHUB_SYSTEM_REFERENCE.docx"
BLUE = "2E74B5"
NAVY = "1F4D78"
PALE = "E8EEF5"
INK = "203746"
MUTED = "667785"


def git(*args: str) -> str:
    return subprocess.check_output(["git", *args], cwd=ROOT, text=True).strip()


def route_path(path: Path) -> str:
    rel = path.relative_to(ROOT / "src" / "app" / "api")
    parts = list(rel.parts[:-1])
    return "/api" + ("/" + "/".join(parts) if parts else "")


def inventory_frontend():
    rows = []
    for p in sorted((ROOT / "src" / "app" / "api").rglob("route.ts")):
        text = p.read_text(encoding="utf-8")
        methods = re.findall(r"export\s+(?:async\s+function|const)\s+(GET|POST|PUT|PATCH|DELETE)", text)
        methods += re.findall(r"export\s*\{\s*(GET|POST|PUT|PATCH|DELETE)\s*\}\s*from", text)
        targets = sorted(set(re.findall(r'["`](/(?:admin|auth|dashboard|finance|health|kyc|ledger|mall|marketplace|matrix|membership|payments|presale|profiles|referrals|registration|rootsbank|routing|shares|subscriptions|theme|vouchers|wallets|whatsapp)[^"`?]*)["`]', text)))
        rows.append((route_path(p), ", ".join(methods), ", ".join(targets) or "Local gateway/session logic", str(p.relative_to(ROOT))))
    return rows


def inventory_encore():
    rows = []
    pat = re.compile(r"export const\s+(\w+)\s*=\s*api(?:\.raw)?", re.M)
    for p in sorted((ROOT / "encore" / "domains").rglob("*.ts")):
        text = p.read_text(encoding="utf-8")
        for m in pat.finditer(text):
            next_export = text.find("export const", m.end())
            chunk = text[m.start():next_export if next_export >= 0 else len(text)]
            method = re.search(r'method:\s*"([A-Z]+)"', chunk)
            path = re.search(r'path:\s*"([^"]+)"', chunk)
            expose = re.search(r'expose:\s*(true|false)', chunk)
            auth = re.search(r'auth:\s*(true|false)', chunk)
            if method and path:
                rows.append((path.group(1), method.group(1), m.group(1), "Public" if expose and expose.group(1) == "true" else "Internal", "Required" if auth and auth.group(1) == "true" else "Endpoint-specific", str(p.relative_to(ROOT))))
    return sorted(rows, key=lambda x: (x[0], x[1], x[2]))


def sql_summary(text: str) -> str:
    items = []
    for verb, name in re.findall(r'(?im)^\s*(CREATE TABLE(?: IF NOT EXISTS)?|ALTER TABLE|CREATE INDEX(?: IF NOT EXISTS)?|CREATE (?:OR REPLACE )?FUNCTION|CREATE TRIGGER|DROP TRIGGER)\s+(?:IF NOT EXISTS\s+)?(?:ONLY\s+)?([\w.]+)', text):
        items.append(f"{verb.upper()} {name}")
    inserts = sorted(set(re.findall(r'(?im)^\s*INSERT INTO\s+([\w.]+)', text)))
    for name in inserts:
        items.append(f"DATA SEED/UPDATE {name}")
    return "; ".join(items) or "Constraint, data repair, or policy change; inspect SQL for exact statement."


def inventory_migrations():
    rows = []
    for p in sorted((ROOT / "encore" / "migrations").rglob("*.sql"), key=lambda x: (x.parent.name, int(x.name.split("_")[0]))):
        text = p.read_text(encoding="utf-8")
        rows.append((p.parent.name, p.name, sql_summary(text), str(p.relative_to(ROOT))))
    return rows


def inventory_tests():
    frontend = list((ROOT / "src").rglob("*.test.ts")) + list((ROOT / "src").rglob("*.test.tsx"))
    backend = list((ROOT / "encore" / "domains").rglob("*.test.ts"))
    return len(frontend), len(backend)


def add_cell_text(cell, text, bold=False, color=None, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW")) or OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    if tc_w.getparent() is None:
        tc_pr.append(tc_w)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "9360"); tbl_w.set(qn("w:type"), "dxa")
    if tbl_w.getparent() is None: tbl_pr.append(tbl_w)
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tbl_pr.append(ind)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; set_cell_width(c, widths[i]); shade(c, PALE); add_cell_text(c, h, True, NAVY, 8)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); tr_pr.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        row_pr = table.rows[-1]._tr.get_or_add_trPr(); row_pr.append(OxmlElement("w:cantSplit"))
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP; add_cell_text(cells[i], value, False, INK, 7.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 14, 7), ("Heading 3", 12, NAVY, 10, 5)]:
        s = styles[name]; s.font.name = "Calibri"; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True
    if "Small Note" not in styles:
        s = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH); s.font.name = "Calibri"; s.font.size = Pt(8); s.font.color.rgb = RGBColor.from_string(MUTED)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25


def add_footer(section, revision):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"KaSiHUB System Reference | Revision {revision} | Controlled engineering document | Page ")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); p._p.append(fld)


def build_doc(frontend, backend, migrations, rev, branch, origin, test_counts):
    doc = Document(); setup_styles(doc)
    sec = doc.sections[0]; sec.top_margin = Inches(.7); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.85); sec.right_margin = Inches(.85)
    add_footer(sec, rev[:7])
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(105); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ENGINEERING SYSTEM REFERENCE"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KaSiHUB"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RGBColor.from_string(INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software scope, architecture, APIs, data model, migrations, integrations and operational controls"); r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()
    add_table(doc, ["Control", "Value"], [("Repository revision", rev), ("Branch", branch), ("Remote main", origin), ("Generated", str(date.today())), ("Evidence basis", "Static repository inspection plus separately recorded deployment verification; no secrets included.")], [2700, 6660])
    doc.add_page_break()

    doc.add_heading("1. Executive system map", 1)
    doc.add_paragraph("KaSiHUB is a multi-domain membership, KYC, payments, network, commerce and share-management platform. The browser-facing Next.js application acts as a controlled gateway. Encore owns the authoritative service contracts and state. PostgreSQL resources are separated by bounded context; presale data is deliberately isolated from the live share ledger until an explicit incorporation workflow runs.")
    add_table(doc, ["Layer", "Scope", "Authority"], [
        ("Web application", "Next.js 16 / React 19 pages, server routes, session cookies, admin and applicant portals", "Presentation and gateway policy; not financial source of truth"),
        ("Backend", "Encore TypeScript domain APIs and service-to-service workflows", "Business rules, authentication/authorization, orchestration"),
        ("Data", "11 PostgreSQL databases, private documents bucket, managed cache", "Database records authoritative; cache is disposable"),
        ("External rails", "Didit KYC, WebPay card checkout, USDT/Remitano evidence flow, email and WhatsApp delivery", "Provider evidence must be authenticated and reconciled"),
        ("Delivery", "GitHub Actions, Vercel frontend, Encore Cloud backend", "Exact revision and environment must be proven separately")], [1500, 4300, 3560])
    doc.add_heading("Status vocabulary", 2)
    for text in ["Implemented: code or migration exists in this revision.", "Tested: automated or manual evidence exists for the stated behavior; implementation alone is not proof.", "Deployed: a named environment has been activated and probed independently.", "Unknown: live provider, infrastructure, data, or operational state was not proven by repository inspection."]:
        add_bullet(doc, text)
    doc.add_heading("Current release evidence", 2)
    add_table(doc, ["Area", "Status", "Evidence / caveat"], [
        ("Repository", "Implemented", f"Local {branch} at {rev}; origin/main at {origin}."),
        ("Backend WebPay rail", "Deployed", "Encore production was activated for this revision; presale migration version 11 was observed clean in production."),
        ("Backend probes", "Tested", "Unauthenticated checkout rejected with 401; malformed WebPay notification rejected with 400."),
        ("Frontend WebPay selector", "Not released", "Local revision is ahead of origin/main; production web UI must not be described as carrying the commit until pushed and deployed."),
        ("Live paid transaction", "Unknown / gated", "No real-money checkout or settlement is claimed by this document."),
    ], [1800, 1500, 6060])

    doc.add_heading("2. Product and domain scope", 1)
    domains = [
        ("Identity", "Registration, login/logout, sessions, profiles, guardians, beneficiaries, roles and admin bootstrap."),
        ("Membership", "Plans, subscriptions, provider payment records and subscription activation."),
        ("KYC", "Local cases, international evidence, Didit sessions/webhooks, review and fail-closed approval policy."),
        ("Payments", "Obligations, receiving configurations, intents, attempts, events, custody evidence and state history."),
        ("Presale", "Invitation-only campaigns, applications, reservations, Remitano/USDT or WebPay selection, evidence, expiry and controlled incorporation."),
        ("Shares", "Sale phases, purchases, certificates, revocation/reissue, inventory evidence and presale incorporation."),
        ("Finance / wallets", "Double-entry ledger, balances, holds, dividends, distributions, reconciliation and operational retries."),
        ("Network", "Matrix placement, ancestry/path records, downline and admin tree."),
        ("Commerce", "Marketplace, RootsBank, mall/NFC transactions and silo configuration."),
        ("Engagement", "Referrals, vouchers, WhatsApp verification/delivery, subscription notices and outbox."),
        ("Administration", "Theme/config versioning, audit logs, reporting, member views and operational control plane."),
    ]
    add_table(doc, ["Domain", "Responsibility"], domains, [1900, 7460])

    doc.add_heading("3. Trust boundaries and critical flows", 1)
    doc.add_heading("Registration and access", 2)
    doc.add_paragraph("The supported web registration path routes citizenship and membership classification through server-authoritative policy, then uses secure registration. Legacy exposed registration remains a bypass risk until retired or hardened. Sessions are scoped, and admin/profile access is enforced through shared guards.")
    doc.add_heading("Presale payment and incorporation", 2)
    for t in ["Invitation and campaign policy determine whether an applicant may enter the flow.", "Application/KYC state precedes order creation; the server owns resumability and next-step authority.", "Step 5 allows the buyer to select Remitano/USDT or WebPay card. The WebPay ZAR price is fixed at R450 per share in the current implementation.", "A submitted hash or provider callback is evidence, not an automatic right to shares; settlement must be authenticated and reconciled.", "Confirmed presale rows remain isolated until an administrator prepares and applies an idempotent incorporation batch into the shares ledger."]:
        add_bullet(doc, t)
    doc.add_heading("External integration boundary", 2)
    add_table(doc, ["Integration", "Purpose", "Secret / control boundary"], [
        ("WebPay / Omnea", "Create hosted card checkout and receive provider notification", "Merchant identifiers and signing credentials are environment secrets; notification validation fails closed."),
        ("Remitano / USDT verifier", "Display receiving instructions and reconcile blockchain settlement evidence", "Receiving configuration is versioned; webhook authentication and chain confirmation remain authoritative."),
        ("Didit", "International identity verification", "Session/webhook identifiers persist; webhook processing is recorded for idempotency/reconciliation."),
        ("Email", "Applicant/reservation communication and retry tracking", "Outcomes are recorded; failed email must not masquerade as success."),
        ("WhatsApp", "Verification codes and voucher delivery", "Attempts, expiry, dedupe and outbox state are persisted."),
    ], [1700, 3700, 3960])

    doc.add_heading("4. Frontend route inventory", 1)
    doc.add_paragraph(f"The Next.js application defines {len(frontend)} route modules. A route may expose more than one HTTP method. 'Backend target' is statically extracted where a literal Encore path is present; local gateway/session logic is reported explicitly.", style="Small Note")
    add_table(doc, ["Public route", "Method(s)", "Backend target / role", "Source"], frontend, [2350, 900, 3810, 2300])

    doc.add_heading("5. Encore endpoint inventory", 1)
    public_count = sum(1 for x in backend if x[3] == "Public")
    internal_count = len(backend) - public_count
    doc.add_paragraph(f"This revision declares {len(backend)} Encore endpoints: {public_count} exposed at the service edge and {internal_count} internal-only. Authentication is summarized from endpoint metadata; endpoint-specific means the handler or shared guard determines access.", style="Small Note")
    add_table(doc, ["Path", "Method", "Handler", "Exposure", "Auth", "Source"], backend, [2800, 700, 1800, 900, 1050, 2110])

    doc.add_heading("6. Data estate", 1)
    dbs = [
        ("identity", "Users, roles, profiles, sessions, registration workflows, admin bootstrap"), ("membership", "Plans, subscriptions, membership payment records"),
        ("network", "Matrix nodes/events and network wallet projections"), ("finance", "Ledger, financial operations, balances, holds, dividends, distributions, reconciliation"),
        ("payments", "Receiving configs, obligations, intents, attempts, events, state history, custody evidence"), ("kyc", "Cases, document metadata, Didit events and reconciliation"),
        ("shares", "Phases, purchases, certificates, holdings and inventory evidence"), ("commerce", "Marketplace, RootsBank, mall transactions, silos"),
        ("engagement", "Referrals, vouchers, contacts, verification, notifications and outbox"), ("audit", "Cross-domain audit log"),
        ("presale", "Campaigns, invitations, applications, encrypted snapshots, orders, payment events, email, incorporation"),
    ]
    add_table(doc, ["Database", "Authoritative contents"], dbs, [1800, 7560])
    add_bullet(doc, "Private object bucket: documents. KYC evidence bytes are private; database rows retain metadata and storage keys.")
    add_bullet(doc, "Managed cache: application-cache with allkeys-lru. Cache records are non-authoritative.")

    doc.add_heading("7. Complete migration register", 1)
    doc.add_paragraph(f"All {len(migrations)} SQL migration files in encore/migrations are listed below. Migrations are ordered numerically within each database. A missing sequence number is preserved as repository history, not silently renumbered.", style="Small Note")
    by_db = defaultdict(list)
    for db, name, summary, source in migrations: by_db[db].append((name, summary, source))
    for db in sorted(by_db):
        doc.add_heading(f"{db} database ({len(by_db[db])} migrations)", 2)
        add_table(doc, ["Migration", "Schema/data action", "Source"], by_db[db], [1900, 5260, 2200])

    doc.add_heading("8. Security, integrity and failure controls", 1)
    for t in [
        "Authentication and authorization are centralized through shared Encore auth/access guards; raw provider endpoints require their own signature and replay controls.",
        "Financial operations carry operation IDs, durable steps, reconciliation runs/findings and idempotency constraints. State transitions leave database evidence.",
        "Payment receiving configuration is versioned and auditable. Optional custody reconciliation fails closed when required capability is absent.",
        "International KYC approval has a database-level policy guard; private document bytes are kept outside SQL in non-public object storage.",
        "Presale applications and sensitive review fields include ciphertext, nonce, authentication-tag and key-version metadata.",
        "No secrets, passwords, webhook keys or provider signing material are included in this reference.",
    ]: add_bullet(doc, t)

    doc.add_heading("9. Verification and delivery", 1)
    add_table(doc, ["Gate", "Repository contract"], [
        ("Frontend", "npm ci; lint; typecheck; coverage; production audit; build; Playwright Chromium"),
        ("Backend", "npm ci; Encore authentication; encore check; encore test; coverage artifact"),
        ("Production", "Controlled full-root deployment plus health and registration policy canaries"),
        ("Observed test inventory", f"{test_counts[0]} frontend test files and {test_counts[1]} Encore domain test files (count only, not a claim that the current run passed)."),
    ], [2100, 7260])
    doc.add_paragraph("Encore staging was retired in August 2026. Backend releases use a full-root production deployment followed by production health and policy checks.")

    doc.add_heading("10. Known gaps and launch blockers", 1)
    gaps = [
        ("Critical", "Frontend release divergence", "Push the exact reviewed commit, let Vercel build it, then prove the rendered Step 5 choice and browser/API behavior."),
        ("High", "Legacy public registration", "Retire, hide or enforce secure server-derived policy on /registration/start."),
        ("High", "Provider production acceptance", "Complete an approved WebPay sandbox/production transaction and authenticated notification test before real customer use."),
        ("High", "Abuse controls", "Add and verify rate limiting/anomaly controls on login, registration and provider-cost endpoints."),
        ("Medium", "Payment FK validation", "Inspect/backfill historical rows and validate the deferred payment intent obligation foreign key in a follow-up migration."),
        ("Governance", "Legal/compliance readiness", "Treat legal approval, privacy, KYC/AML policy and financial operations ownership as separate release gates."),
    ]
    add_table(doc, ["Severity", "Gap", "Required closure"], gaps, [1200, 2700, 5460])

    doc.add_heading("11. Ownership and maintenance rules", 1)
    for t in ["Update this document on any added/removed endpoint, database, migration, provider rail or deployment boundary.", "Never edit an applied migration; add a new numbered migration and verify production schema_migrations state.", "Pin releases to a Git revision and record frontend and backend deployment evidence separately.", "Do not infer a live feature from source presence, a commit, a successful build or a dashboard deployment banner.", "Keep provider secrets in environment secret stores and never place them in source."]:
        add_bullet(doc, t)
    doc.add_paragraph("End of controlled reference.", style="Small Note")
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def build_md(frontend, backend, migrations, rev, branch, origin):
    lines = ["# KaSiHUB System Reference", "", f"Generated: {date.today()}  ", f"Revision: `{rev}`  ", f"Branch: `{branch}`  ", f"origin/main: `{origin}`", "", "> This source appendix is generated from the repository. The DOCX is the controlled reading copy.", "", "## Inventory totals", "", f"- Next.js route modules: {len(frontend)}", f"- Encore endpoints: {len(backend)}", f"- SQL migrations: {len(migrations)}", f"- SQL databases: {len(set(x[0] for x in migrations))}", "", "## Frontend routes", "", "| Route | Methods | Backend target / role | Source |", "|---|---|---|---|"]
    lines += [f"| `{a}` | {b} | `{c}` | `{d}` |" for a,b,c,d in frontend]
    lines += ["", "## Encore endpoints", "", "| Path | Method | Handler | Exposure | Auth | Source |", "|---|---|---|---|---|---|"]
    lines += [f"| `{a}` | {b} | `{c}` | {d} | {e} | `{f}` |" for a,b,c,d,e,f in backend]
    lines += ["", "## Migrations", "", "| Database | Migration | Action | Source |", "|---|---|---|---|"]
    lines += [f"| `{a}` | `{b}` | {c.replace('|','/')} | `{d}` |" for a,b,c,d in migrations]
    OUT.mkdir(parents=True, exist_ok=True); MD_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    frontend = inventory_frontend(); backend = inventory_encore(); migrations = inventory_migrations()
    rev = git("rev-parse", "HEAD"); branch = git("branch", "--show-current"); origin = git("rev-parse", "origin/main")
    build_md(frontend, backend, migrations, rev, branch, origin)
    build_doc(frontend, backend, migrations, rev, branch, origin, inventory_tests())
    print(MD_PATH); print(DOCX_PATH)
